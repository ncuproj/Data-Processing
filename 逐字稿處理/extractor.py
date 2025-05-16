import sys
import io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

import os
import re
from collections import defaultdict
from docx import Document

# 設定主資料夾
root_folder = ".\議會速紀錄"
output_folder = ".\議員分類"

# 確保輸出資料夾存在
os.makedirs(output_folder, exist_ok=True)

# 正則表達式
date_pattern = re.compile(r"(\d{3})\s*年\s*(\d{1,2})\s*月\s*(\d{1,2})\s*日")
speech_pattern = re.compile(r"([一-龥]{1,2}議員[一-龥]{1,2})\s*[:：]")  # 只匹配"X議員XX："
non_speaker_pattern = re.compile(r"([一-龥])?(市長|主席)[一-龥]*\s*[:：]")  # 過濾非議員發言

# 不合法的檔名字元
INVALID_CHARS = r'[<>:"/\\|?*]'

def sanitize_filename(name):
    """ 移除 Windows 不合法字元 """
    return re.sub(INVALID_CHARS, "", name)

def extract_text_from_docx(docx_path):
    """ 讀取 docx 檔案並回傳所有文字段落 """
    doc = Document(docx_path)
    return [para.text.strip() for para in doc.paragraphs if para.text.strip()]

def save_speech_to_docx(member_name, date, speech):
    """ 儲存議員發言到 docx """
    if not date:
        date = "未知日期"

    safe_member_name = sanitize_filename(member_name)
    member_folder = os.path.join(output_folder, safe_member_name)
    os.makedirs(member_folder, exist_ok=True)

    docx_filename = f"{date}.docx"
    docx_path = os.path.join(member_folder, docx_filename)

    if os.path.exists(docx_path):
        doc = Document(docx_path)
    else:
        doc = Document()

    doc.add_paragraph(speech)
    doc.save(docx_path)

def get_docx_groups_by_folder(directory):
    """ 遞迴獲取所有子資料夾內的 .docx 檔案，並按「去掉 `-X` 的檔名」在同一個資料夾內分類 """
    folder_docx_map = {}

    for root, _, files in os.walk(directory):
        docx_groups = defaultdict(list)
        
        for file in files:
            if file.lower().endswith(".docx"):
                base_name = re.sub(r"-\d+", "", file)  # 去掉 `-X`
                full_path = os.path.join(root, file)
                docx_groups[base_name].append(full_path)

        if docx_groups:
            folder_docx_map[root] = docx_groups  # 只記錄有 `.docx` 的資料夾

    return folder_docx_map

# 處理每個子資料夾
all_folders = get_docx_groups_by_folder(root_folder)

for folder_path, docx_groups in all_folders.items():
    for base_name, docx_files in docx_groups.items():
        merged_text = []
        first_docx_date = None

        # 先尋找日期
        for docx_file in docx_files:
            paragraphs = extract_text_from_docx(docx_file)
            for para in paragraphs:
                date_match = date_pattern.search(para)
                if date_match:
                    year, month, day = date_match.groups()
                    first_docx_date = f"{year}{month.zfill(2)}{day.zfill(2)}"
                    break
            if first_docx_date:
                break

        if not first_docx_date:
            first_docx_date = "未知日期"

        # 合併相同檔名的 .docx
        for docx_file in docx_files:
            merged_text.extend(extract_text_from_docx(docx_file))

        # 儲存合併後的 .docx
        merged_docx_path = os.path.join(folder_path, f"{base_name}_合併.docx")
        doc = Document()
        for para in merged_text:
            doc.add_paragraph(para)
        doc.save(merged_docx_path)

        print(f"已合併 {base_name} 的 .docx，存為 {merged_docx_path}")

        # 擷取議員發言
        paragraphs = extract_text_from_docx(merged_docx_path)

        current_speaker = None
        speech_content = []

        for para in paragraphs:
            # 遇到"……"則視為分隔
            if "……" in para:
                para = para.replace("……", "。")  # 避免影響發言判斷

            # 檢測是否為議員發言
            speaker_match = speech_pattern.match(para)
            non_speaker_match = non_speaker_pattern.match(para)

            if speaker_match:
                speaker_name = speaker_match.group(1)

                # 儲存前一位議員的發言
                if current_speaker and speech_content:
                    save_speech_to_docx(current_speaker, first_docx_date, "\n".join(speech_content))

                # 更新當前發言者
                current_speaker = speaker_name
                speech_content = [para]

            elif current_speaker:
                # 檢測是否遇到非議員發言，若遇到則終止
                if non_speaker_match:
                    if speech_content:
                        save_speech_to_docx(current_speaker, first_docx_date, "\n".join(speech_content))
                    current_speaker = None
                    speech_content = []
                else:
                    speech_content.append(para)

        # 儲存最後一位議員的發言
        if current_speaker and speech_content:
            save_speech_to_docx(current_speaker, first_docx_date, "\n".join(speech_content))

print("所有議員發言已成功分類並儲存！")
