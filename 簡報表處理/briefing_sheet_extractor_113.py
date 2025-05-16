import sys
import io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

import re
from pathlib import Path
from docx import Document
from openpyxl import Workbook, load_workbook
from glob import glob
import tkinter as tk
from tkinter import filedialog, messagebox, scrolledtext

# --- 核心邏輯函式區 ---
# 從文字中提取日期
def extract_latest_date(text):
    match = re.search(r"中華民國(\d{3})年(\d{1,2})月(\d{1,2})日", text)
    re_match = re.search(r"中華民國(\d{3})年(\d{1,2})月(\d{1,2})", text)
    if not match:
        if not re_match:
            return None
        else:
            year, month, day = re_match.groups()
    else:
        year, month, day = match.groups()
    return f"{year}{int(month):02d}{int(day):02d}"

# 只保留中文字
def clean_chinese_name(name):
    return re.sub(r"[^\u4e00-\u9fff]", "", name)

# 處理局處室
def extract_agencies(topic_text):
    agencies = []
    if re.search(r"擬不分辦|無涉機關業務", topic_text):
        return "其他"
    match = re.search(r"（([^）]*)）", topic_text)
    if match:
        content = match.group(1)
        parts = re.split(r'[；;]', content)
        for part in parts:
            sub_match = re.search(r"[主協]：(.+)", part)
            if sub_match:
                raw_agencies = sub_match.group(1)
                split_agencies = re.split(r'[、,，]', raw_agencies)
                agencies.extend(clean_chinese_name(a) for a in split_agencies if a)
    return ";".join(filter(None, agencies)) if agencies else "其他"

# 逐筆擷取資訊（包含本地日期）
def extract_info_from_docx(docx_path):
    doc = Document(docx_path)
    full_text = "\n".join(p.text for p in doc.paragraphs)
    full_text += "\n" + "\n".join(cell.text for table in doc.tables for row in table.rows for cell in row.cells)
    global_date = extract_latest_date(full_text) or "unknown"
    records = []

    for table in doc.tables:
        member_name = None
        for row in table.rows:
            row_text = " ".join(cell.text for cell in row.cells)
            row_date = extract_latest_date(row_text)
            if row_date != None:
                effective_date = row_date or global_date
            cells = row.cells
            for idx, cell in enumerate(cells):
                text = cell.text.strip()
                if "議員姓名" in text and idx + 1 < len(cells):
                    raw_name = cells[idx + 1].text.strip()
                    member_name = clean_chinese_name(raw_name)
                elif re.search(r"質詢\s*議題", text) and idx + 1 < len(cells):
                    topic = cells[idx + 1].text.strip()
                    if topic and not re.search(r"質詢\s*議題", topic) and member_name:
                        topic_cleaned = re.sub(r"\s+", "", topic)
                        agencies = extract_agencies(topic_cleaned)
                        records.append((member_name, topic_cleaned, agencies, effective_date))
    return list(dict.fromkeys(records))  # 去重

# 儲存資料到 Excel
def save_to_excel(member, topic, agencies, date_str, base_dir="output"):
    folder = Path(base_dir) / member
    folder.mkdir(parents=True, exist_ok=True)
    file_path = folder / f"{date_str}.xlsx"
    if file_path.exists():
        wb = load_workbook(file_path)
        ws = wb.active
    else:
        wb = Workbook()
        ws = wb.active
        ws.append(["質詢議題", "局處室"])
    ws.append([topic, agencies])
    wb.save(file_path)

# --- UI 功能區 ---

def select_files():
    file_paths = filedialog.askopenfilenames(title="選擇 Word 檔案", filetypes=[("Word files", "*.docx")])
    if file_paths:
        file_list.set("\n".join(file_paths))

def process_files():
    paths = file_list.get().strip().splitlines()
    if not paths:
        messagebox.showwarning("未選擇檔案", "請先選擇 Word 檔案")
        return

    output_text.delete("1.0", tk.END)
    for path in paths:
        output_text.insert(tk.END, f"處理中：{path}\n")
        try:
            records = extract_info_from_docx(path)
            output_text.insert(tk.END, f"共抓到 {len(records)} 筆資料\n")
            for member, topic, agencies, date_str in records:
                output_text.insert(tk.END, f"{member} | {date_str} | {topic[:20]}... | {agencies}\n")
                save_to_excel(member, topic, agencies, date_str)
        except Exception as e:
            output_text.insert(tk.END, f"錯誤：{path} - {e}\n")

    messagebox.showinfo("完成", "所有檔案處理完畢！")

# --- GUI 組件初始化 ---

root = tk.Tk()
root.title("市政總質詢處理工具")
root.geometry("700x500")

frame = tk.Frame(root)
frame.pack(pady=10)

tk.Button(frame, text="選擇 Word 檔案", command=select_files).grid(row=0, column=0, padx=5)
tk.Button(frame, text="開始處理", command=process_files).grid(row=0, column=1, padx=5)

file_list = tk.StringVar()
tk.Label(root, text="選擇的檔案：").pack(anchor="w", padx=10)
tk.Entry(root, textvariable=file_list, width=100).pack(padx=10)

output_text = scrolledtext.ScrolledText(root, wrap=tk.WORD, height=20)
output_text.pack(padx=10, pady=10, fill=tk.BOTH, expand=True)

root.mainloop()
